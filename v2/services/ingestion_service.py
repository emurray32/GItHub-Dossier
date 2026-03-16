"""
Ingestion Service — Smart ingestion layer that turns raw data into
actionable BDR intelligence.

Accepts any format (CSV, Excel, DOCX), enriches via Apollo, evaluates
each signal from a BDR perspective, and deduplicates across uploads.

Every ingestion path follows the same pattern:
  1. Parse the file (structured or unstructured)
  2. Find or create accounts (with dedup + enrichment)
  3. Create intent signals
  4. Post-process: Apollo enrichment + LLM BDR evaluation
  5. Log activity

Errors are captured per-row so one bad record never crashes the batch.
"""
import csv
import io
import json
import logging
import os
import re
import uuid
from typing import Optional

from v2.db import db_connection, row_to_dict, rows_to_dicts, safe_json_dumps
from v2.services import activity_service
from v2.services import campaign_service
from v2.services import signal_service
from v2.services import account_service

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# LLM Client — shared module (Gemini Flash primary, OpenAI fallback)
# ---------------------------------------------------------------------------

from v2.services.llm_client import llm_generate as _llm_generate, get_llm_client as _get_llm_client


# ---------------------------------------------------------------------------
# CSV Ingestion
# ---------------------------------------------------------------------------

def ingest_csv(
    file_content: bytes,
    source_label: str = 'csv_upload',
    created_by: Optional[str] = None,
) -> dict:
    """Parse a CSV and create one intent signal per valid row.

    Required columns: company_name, signal_description
    Optional columns: website, signal_type, evidence, industry,
                      company_size, annual_revenue, account_owner

    Returns:
        dict with keys: signals_created, accounts_created, accounts_matched,
                        errors (list of strings), batch_id
    """
    batch_id = uuid.uuid4().hex[:12]
    result = {
        'signals_created': 0,
        'accounts_created': 0,
        'accounts_matched': 0,
        'errors': [],
        'batch_id': batch_id,
    }

    # --- Decode ---
    try:
        text = file_content.decode('utf-8-sig')
    except UnicodeDecodeError:
        try:
            text = file_content.decode('latin-1')
        except UnicodeDecodeError:
            result['errors'].append('File encoding not supported (use UTF-8)')
            return result

    # --- Parse CSV ---
    reader = csv.DictReader(io.StringIO(text))
    headers = reader.fieldnames or []
    headers_lower = [h.lower().strip() for h in headers]

    # Map flexible header names to canonical keys
    company_col = _find_column(headers, headers_lower,
                               ('company_name', 'company', 'name', 'account_name'))
    signal_col = _find_column(headers, headers_lower,
                              ('signal_description', 'signal', 'description'))

    if not company_col:
        result['errors'].append(
            'CSV must have a company_name column '
            '(also accepts: company, name, account_name)')
        return result

    if not signal_col:
        result['errors'].append(
            'CSV must have a signal_description column '
            '(also accepts: signal, description)')
        return result

    # Optional columns
    website_col = _find_column(headers, headers_lower,
                               ('website', 'domain', 'website_url', 'url'))
    signal_type_col = _find_column(headers, headers_lower,
                                   ('signal_type', 'type'))
    evidence_col = _find_column(headers, headers_lower,
                                ('evidence', 'evidence_value'))
    industry_col = _find_column(headers, headers_lower,
                                ('industry',))
    size_col = _find_column(headers, headers_lower,
                            ('company_size', 'size', 'employees'))
    revenue_col = _find_column(headers, headers_lower,
                               ('annual_revenue', 'revenue'))
    owner_col = _find_column(headers, headers_lower,
                             ('account_owner', 'owner'))

    # --- Process rows ---
    new_account_ids = []
    created_signal_data = []

    for row_num, row in enumerate(reader, start=2):
        try:
            company_name = (row.get(company_col) or '').strip()
            signal_desc = (row.get(signal_col) or '').strip()

            if not company_name:
                result['errors'].append(f'Row {row_num}: missing company_name')
                continue
            if not signal_desc:
                result['errors'].append(f'Row {row_num}: missing signal_description')
                continue

            # Extract optional fields
            website = (row.get(website_col) or '').strip() if website_col else None
            signal_type = (row.get(signal_type_col) or '').strip() if signal_type_col else None
            evidence = (row.get(evidence_col) or '').strip() if evidence_col else None
            industry = (row.get(industry_col) or '').strip() if industry_col else None
            company_size = (row.get(size_col) or '').strip() if size_col else None
            annual_revenue = (row.get(revenue_col) or '').strip() if revenue_col else None
            account_owner = (row.get(owner_col) or '').strip() if owner_col else None

            # 1. Find or create account
            existing = account_service.find_account_by_name(company_name)
            if existing:
                account_id = existing['id']
                result['accounts_matched'] += 1
                # Enrich existing account with any new data from this upload
                account_service.update_account_enrichment(
                    account_id,
                    website=website, industry=industry,
                    company_size=company_size, annual_revenue=annual_revenue,
                )
            else:
                account_id = account_service.find_or_create_account(
                    company_name=company_name,
                    website=website or None,
                    industry=industry or None,
                    company_size=company_size or None,
                    annual_revenue=annual_revenue or None,
                    account_owner=account_owner or None,
                )
                result['accounts_created'] += 1
                new_account_ids.append({
                    'account_id': account_id,
                    'company_name': company_name,
                    'website': website,
                })

            # 2. Auto-recommend campaign
            rec = campaign_service.recommend_campaign(
                signal_type=signal_type or None,
            )

            # 3. Create intent signal
            signal_id = signal_service.create_signal(
                account_id=account_id,
                signal_description=signal_desc,
                signal_type=signal_type or None,
                evidence_type='csv_import',
                evidence_value=evidence or None,
                signal_source='csv_upload',
                recommended_campaign_id=rec.get('campaign_id'),
                recommended_campaign_reasoning=rec.get('reasoning'),
                created_by=created_by,
                ingestion_batch_id=batch_id,
                raw_payload=safe_json_dumps(dict(row)),
            )

            result['signals_created'] += 1
            created_signal_data.append({
                'signal_id': signal_id,
                'company_name': company_name,
                'signal_description': signal_desc,
                'signal_type': signal_type,
            })

            # 4. Log activity
            activity_service.log_activity(
                event_type='signal_created',
                entity_type='signal',
                entity_id=signal_id,
                details={
                    'source': 'csv_upload',
                    'batch_id': batch_id,
                    'company_name': company_name,
                    'signal_type': signal_type,
                },
                created_by=created_by,
            )

        except Exception as exc:
            logger.exception("[INGEST] Error on CSV row %d", row_num)
            result['errors'].append(f'Row {row_num}: {str(exc)[:200]}')

    # Log batch-level activity
    activity_service.log_activity(
        event_type='csv_imported',
        entity_type='batch',
        entity_id=None,
        details={
            'batch_id': batch_id,
            'source_label': source_label,
            'signals_created': result['signals_created'],
            'accounts_created': result['accounts_created'],
            'accounts_matched': result['accounts_matched'],
            'error_count': len(result['errors']),
        },
        created_by=created_by,
    )

    logger.info(
        "[INGEST] CSV batch %s complete: %d signals, %d new accounts, %d matched, %d errors",
        batch_id, result['signals_created'], result['accounts_created'],
        result['accounts_matched'], len(result['errors']),
    )

    # --- Post-processing: Apollo enrichment + BDR evaluation ---
    enrichment = _post_process_batch(new_account_ids, created_signal_data)
    result['enrichment'] = enrichment

    return result


# ---------------------------------------------------------------------------
# Manual / Single-Signal Ingestion
# ---------------------------------------------------------------------------

def ingest_manual(
    account_id: int,
    signal_description: str,
    signal_type: Optional[str] = None,
    evidence_value: Optional[str] = None,
    created_by: Optional[str] = None,
) -> int:
    """Create a single intent signal for an existing account.

    Returns:
        The new signal id.
    """
    # Auto-recommend campaign
    rec = campaign_service.recommend_campaign(signal_type=signal_type)

    signal_id = signal_service.create_signal(
        account_id=account_id,
        signal_description=signal_description,
        signal_type=signal_type,
        evidence_type='manual',
        evidence_value=evidence_value,
        signal_source='manual_entry',
        recommended_campaign_id=rec.get('campaign_id'),
        recommended_campaign_reasoning=rec.get('reasoning'),
        created_by=created_by,
    )

    activity_service.log_activity(
        event_type='signal_created',
        entity_type='signal',
        entity_id=signal_id,
        details={
            'source': 'manual_entry',
            'signal_type': signal_type,
            'account_id': account_id,
        },
        created_by=created_by,
    )

    return signal_id


# ---------------------------------------------------------------------------
# Excel / Smart File Ingestion
# ---------------------------------------------------------------------------

# Canonical field -> list of header substrings that match (checked via 'in', case-insensitive).
# Order matters: first match wins.  Entries are (canonical_key, match_phrases).
_COLUMN_MATCHERS = [
    ('company_name',        ['company_name', 'company', 'account_name', 'account', 'name']),
    ('signal_description',  ['signal detail', 'signal_description', 'description', 'detail']),
    ('website',             ['domain', 'website_url', 'website', 'url']),
    ('signal_type',         ['signal type', 'signal_type', 'type']),
    ('evidence_value',      ['source url', 'source_url', 'evidence_value', 'evidence']),
    ('company_size',        ['estimated size', 'company_size', 'size', 'employees']),
    ('industry',            ['industry', 'sector', 'vertical']),
    ('account_owner',       ['account_owner', 'owner', 'rep', 'assigned']),
    ('score',               ['score', 'priority', 'weight', 'rating']),
    ('outreach_angle',      ['outreach angle', 'outreach_angle', 'outreach', 'angle']),
    ('status',              ['status']),
    ('notes',               ['notes', 'comment']),
    ('date_found',          ['date found', 'date_found', 'date', 'created']),
    ('buyer_persona',       ['buyer persona', 'buyer_persona', 'persona']),
    ('video_url',           ['video content url', 'video_url', 'video url', 'video']),
    ('annual_revenue',      ['annual_revenue', 'revenue']),
]

# Headers to skip when matching (too generic or ambiguous when alone)
_SKIP_EXACT = {'none', ''}


def _smart_match_columns(headers):
    """Fuzzy-match spreadsheet headers to canonical field names.

    Returns a dict of {canonical_key: original_header_name}.
    Unmatched headers are ignored (their data lands in raw_payload).
    """
    mapping = {}
    used_headers = set()
    headers_lower = [(h or '').strip().lower() for h in headers]

    for canonical, phrases in _COLUMN_MATCHERS:
        for phrase in phrases:
            for idx, low in enumerate(headers_lower):
                if low in _SKIP_EXACT or idx in used_headers:
                    continue
                # "source url" should not match "website" (avoid the generic 'url' grabbing it)
                if low == phrase or phrase in low:
                    mapping[canonical] = headers[idx]
                    used_headers.add(idx)
                    break
            if canonical in mapping:
                break

    return mapping


def _normalize_signal_type(raw_type):
    """Convert human-readable signal type to snake_case identifier.

    'Hiring - Localization' -> 'hiring_localization'
    'YouTube Channel + Academy' -> 'youtube_channel_academy'
    'Hiring - Hidden Role (i18n)' -> 'hiring_hidden_role_i18n'
    """
    if not raw_type:
        return None
    # Remove parens but keep their content: "(i18n)" -> " i18n"
    cleaned = re.sub(r'[()]', ' ', str(raw_type)).strip().lower()
    cleaned = re.sub(r'[\s\-/+]+', '_', cleaned)
    cleaned = re.sub(r'[^a-z0-9_]', '', cleaned)
    cleaned = re.sub(r'_+', '_', cleaned).strip('_')
    return cleaned or None


def _coerce_str(val):
    """Convert a cell value to a trimmed string (handles dates, numbers, None)."""
    if val is None:
        return ''
    if hasattr(val, 'isoformat'):
        return val.isoformat()
    return str(val).strip()


def _process_rows(rows, source_label, created_by, sheet_name=None):
    """Process a list of dicts (one per row) into intent signals.

    This is the shared core logic used by CSV, Excel, and DOCX ingestion.
    Each dict should have canonical keys from _smart_match_columns.

    Returns a result dict with signals_created, accounts_created, etc.
    """
    batch_id = uuid.uuid4().hex[:12]
    result = {
        'signals_created': 0,
        'accounts_created': 0,
        'accounts_matched': 0,
        'skipped': 0,
        'skipped_duplicates': 0,
        'errors': [],
        'batch_id': batch_id,
        'sheet_name': sheet_name,
    }

    new_account_ids = []
    created_signal_data = []

    for row_num, row in enumerate(rows, start=2):
        try:
            company_name = _coerce_str(row.get('company_name'))
            signal_desc = _coerce_str(row.get('signal_description'))

            if not company_name:
                # Skip silently — likely an empty row
                result['skipped'] += 1
                continue
            if not signal_desc:
                result['errors'].append(f'Row {row_num}: missing signal description')
                continue

            website = _coerce_str(row.get('website')) or None
            raw_signal_type = _coerce_str(row.get('signal_type')) or None
            signal_type = raw_signal_type  # Keep human-readable type as-is
            evidence = _coerce_str(row.get('evidence_value')) or None
            industry = _coerce_str(row.get('industry')) or None
            company_size = _coerce_str(row.get('company_size')) or None
            annual_revenue = _coerce_str(row.get('annual_revenue')) or None
            account_owner = _coerce_str(row.get('account_owner')) or None
            outreach_angle = _coerce_str(row.get('outreach_angle')) or None

            # 1. Find or create account
            existing = account_service.find_account_by_name(company_name)
            if existing:
                account_id = existing['id']
                result['accounts_matched'] += 1
                # Enrich existing account with any new data from this upload
                account_service.update_account_enrichment(
                    account_id,
                    website=website, industry=industry,
                    company_size=company_size, annual_revenue=annual_revenue,
                )
            else:
                account_id = account_service.find_or_create_account(
                    company_name=company_name,
                    website=website,
                    industry=industry,
                    company_size=company_size,
                    annual_revenue=annual_revenue,
                    account_owner=account_owner,
                )
                result['accounts_created'] += 1
                new_account_ids.append({
                    'account_id': account_id,
                    'company_name': company_name,
                    'website': website,
                })

            # 2. Check for duplicate signal (match on raw type stored in DB)
            if signal_service.check_duplicate_signal(account_id, signal_type, source_label, evidence_value=evidence):
                result['skipped_duplicates'] += 1
                continue

            # 3. Auto-recommend campaign (normalize for keyword matching)
            normalized_type = _normalize_signal_type(signal_type)
            rec = campaign_service.recommend_campaign(signal_type=normalized_type)

            # 4. Build raw_payload with ALL original data (preserves unmapped fields)
            raw_payload = {k: _coerce_str(v) for k, v in row.items() if v is not None and _coerce_str(v)}

            # 5. Create intent signal
            signal_id = signal_service.create_signal(
                account_id=account_id,
                signal_description=signal_desc,
                signal_type=signal_type,
                evidence_type='file_import',
                evidence_value=evidence,
                signal_source=source_label,
                recommended_campaign_id=rec.get('campaign_id'),
                recommended_campaign_reasoning=rec.get('reasoning'),
                created_by=created_by,
                ingestion_batch_id=batch_id,
                raw_payload=safe_json_dumps(raw_payload),
                outreach_angle=outreach_angle,
            )

            result['signals_created'] += 1
            created_signal_data.append({
                'signal_id': signal_id,
                'company_name': company_name,
                'signal_description': signal_desc,
                'signal_type': signal_type,
            })

            # 6. Log activity
            activity_service.log_activity(
                event_type='signal_created',
                entity_type='signal',
                entity_id=signal_id,
                details={
                    'source': source_label,
                    'batch_id': batch_id,
                    'company_name': company_name,
                    'signal_type': signal_type,
                    'sheet_name': sheet_name,
                },
                created_by=created_by,
            )

        except Exception as exc:
            logger.exception("[INGEST] Error on row %d", row_num)
            result['errors'].append(f'Row {row_num}: {str(exc)[:200]}')

    # Log batch-level activity
    activity_service.log_activity(
        event_type='file_imported',
        entity_type='batch',
        entity_id=None,
        details={
            'batch_id': batch_id,
            'source_label': source_label,
            'sheet_name': sheet_name,
            'signals_created': result['signals_created'],
            'accounts_created': result['accounts_created'],
            'accounts_matched': result['accounts_matched'],
            'skipped': result['skipped'],
            'error_count': len(result['errors']),
        },
        created_by=created_by,
    )

    # --- Post-processing: Apollo enrichment + BDR evaluation ---
    enrichment = _post_process_batch(new_account_ids, created_signal_data)
    result['enrichment'] = enrichment

    return result


def ingest_excel(file_content, source_label='excel_upload', created_by=None, clear_existing=False):
    """Parse an Excel workbook and create intent signals from all valid sheets.

    Args:
        clear_existing: if True, delete all intent_signals and monitored_accounts
                        before importing.

    Returns:
        dict with keys: sheets (list of per-sheet results), totals.
    """
    import openpyxl

    cleared_count = 0
    if clear_existing:
        cleared_count = _clear_all_signals()

    wb = openpyxl.load_workbook(
        filename=io.BytesIO(file_content),
        data_only=True,
        read_only=True,
    )

    all_results = []
    totals = {'signals_created': 0, 'accounts_created': 0, 'accounts_matched': 0,
              'skipped': 0, 'errors': [], 'cleared': cleared_count}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Read all rows into memory (read_only worksheets are generators)
        raw_rows = list(ws.iter_rows(values_only=True))
        if len(raw_rows) < 2:
            continue  # Need at least header + 1 data row

        # First non-empty row is the header
        header_row = raw_rows[0]
        headers = [str(h).strip() if h else '' for h in header_row]

        # Skip sheets with no recognizable headers
        col_map = _smart_match_columns(headers)
        if 'company_name' not in col_map:
            logger.info("[INGEST] Sheet '%s' skipped — no company column found in headers: %s",
                        sheet_name, headers)
            continue

        # Build canonical dicts from data rows
        canonical_rows = []
        for raw_row in raw_rows[1:]:
            row_dict = {}
            for canonical_key, original_header in col_map.items():
                col_idx = headers.index(original_header)
                row_dict[canonical_key] = raw_row[col_idx] if col_idx < len(raw_row) else None

            # Also capture ALL columns for raw_payload
            for idx, hdr in enumerate(headers):
                if hdr and idx < len(raw_row):
                    safe_key = hdr.strip().lower().replace(' ', '_')
                    if safe_key not in row_dict:
                        row_dict[safe_key] = raw_row[idx]

            canonical_rows.append(row_dict)

        if not canonical_rows:
            continue

        # Count non-empty rows (at least company_name has a value)
        real_rows = [r for r in canonical_rows if _coerce_str(r.get('company_name'))]
        if not real_rows:
            logger.info("[INGEST] Sheet '%s' skipped — all rows empty", sheet_name)
            continue

        logger.info("[INGEST] Processing sheet '%s': %d data rows, columns mapped: %s",
                    sheet_name, len(real_rows), list(col_map.keys()))

        sheet_result = _process_rows(
            real_rows,
            source_label=source_label,
            created_by=created_by,
            sheet_name=sheet_name,
        )
        all_results.append(sheet_result)

        # Accumulate totals
        for key in ('signals_created', 'accounts_created', 'accounts_matched', 'skipped'):
            totals[key] += sheet_result.get(key, 0)
        totals['errors'].extend(sheet_result.get('errors', []))

    wb.close()

    return {
        'sheets': all_results,
        'totals': totals,
        'sheets_processed': len(all_results),
        'sheets_total': len(wb.sheetnames),
    }


# ---------------------------------------------------------------------------
# DOCX Ingestion — Structured tables or free-text via LLM
# ---------------------------------------------------------------------------

def ingest_docx(file_content, source_label='docx_upload', created_by=None):
    """Parse a Word document for intent signals.

    Handles two formats:
    1. Tables — treated like spreadsheet rows, uses _smart_match_columns
    2. Free text — sent to LLM for structured extraction

    Returns same shape as _process_rows result, or an error dict.
    """
    try:
        import docx
    except ImportError:
        return {'signals_created': 0, 'errors': ['python-docx not installed']}

    try:
        doc = docx.Document(io.BytesIO(file_content))
    except Exception as e:
        return {'signals_created': 0, 'errors': [f'Could not read DOCX file: {str(e)[:200]}']}

    # 1. Try to extract from tables first (structured data)
    table_rows = _extract_docx_tables(doc)
    if table_rows:
        logger.info("[INGEST] DOCX: found %d rows in tables", len(table_rows))
        return _process_rows(table_rows, source_label, created_by, sheet_name='table')

    # 2. No structured tables — extract all text and use LLM to parse
    full_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    if not full_text.strip():
        return {'signals_created': 0, 'errors': ['Document is empty']}

    logger.info("[INGEST] DOCX: no tables found, using LLM to parse %d chars of text",
                len(full_text))

    parsed_rows = _llm_extract_signals(full_text)
    if not parsed_rows:
        return {
            'signals_created': 0,
            'errors': ['Could not extract signals from document. '
                       'Try a structured format (table, CSV, or Excel).'],
        }

    logger.info("[INGEST] DOCX: LLM extracted %d signals from free text", len(parsed_rows))
    return _process_rows(parsed_rows, source_label, created_by, sheet_name='document')


# ---------------------------------------------------------------------------
# Plain Text Ingestion — Free-text via LLM
# ---------------------------------------------------------------------------

def ingest_text(file_content, source_label='text_upload', created_by=None):
    """Parse a plain text file for intent signals using LLM extraction.

    Returns same shape as _process_rows result, or an error dict.
    """
    try:
        text = file_content.decode('utf-8-sig')
    except UnicodeDecodeError:
        try:
            text = file_content.decode('latin-1')
        except UnicodeDecodeError:
            return {'signals_created': 0, 'errors': ['File encoding not supported (use UTF-8)']}

    text = text.strip()
    if not text:
        return {'signals_created': 0, 'errors': ['File is empty']}

    logger.info("[INGEST] TXT: using LLM to parse %d chars of text", len(text))

    parsed_rows = _llm_extract_signals(text)
    if not parsed_rows:
        return {
            'signals_created': 0,
            'errors': ['Could not extract signals from text file. '
                       'Make sure it contains company names and signal descriptions.'],
        }

    logger.info("[INGEST] TXT: LLM extracted %d signals", len(parsed_rows))
    return _process_rows(parsed_rows, source_label, created_by, sheet_name='text')


# ---------------------------------------------------------------------------
# PDF Ingestion — Extract text and parse via LLM
# ---------------------------------------------------------------------------

def ingest_pdf(file_content, source_label='pdf_upload', created_by=None):
    """Parse a PDF file for intent signals.

    Extracts text from all pages, then uses LLM to identify signals.
    Returns same shape as _process_rows result, or an error dict.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {'signals_created': 0, 'errors': ['PyMuPDF not installed — PDF parsing unavailable']}

    try:
        doc = fitz.open(stream=file_content, filetype='pdf')
    except Exception as e:
        return {'signals_created': 0, 'errors': [f'Could not read PDF: {str(e)[:200]}']}

    # Extract text from all pages
    pages_text = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages_text.append(text.strip())
    doc.close()

    full_text = '\n\n'.join(pages_text)
    if not full_text.strip():
        return {'signals_created': 0, 'errors': ['PDF contains no extractable text']}

    logger.info("[INGEST] PDF: extracted %d chars from %d pages", len(full_text), len(pages_text))

    parsed_rows = _llm_extract_signals(full_text)
    if not parsed_rows:
        return {
            'signals_created': 0,
            'errors': ['Could not extract signals from PDF. '
                       'Make sure it contains company names and signal descriptions.'],
        }

    logger.info("[INGEST] PDF: LLM extracted %d signals", len(parsed_rows))
    return _process_rows(parsed_rows, source_label, created_by, sheet_name='pdf')


def _extract_docx_tables(doc):
    """Extract structured data from DOCX tables using _smart_match_columns."""
    all_rows = []
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        col_map = _smart_match_columns(headers)
        if 'company_name' not in col_map:
            continue

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            row_dict = {}
            for canonical_key, original_header in col_map.items():
                col_idx = headers.index(original_header)
                row_dict[canonical_key] = cells[col_idx] if col_idx < len(cells) else None

            # Capture unmapped columns in raw_payload
            for idx, hdr in enumerate(headers):
                if hdr and idx < len(cells):
                    safe_key = hdr.strip().lower().replace(' ', '_')
                    if safe_key not in row_dict:
                        row_dict[safe_key] = cells[idx]

            all_rows.append(row_dict)

    # Filter out empty rows
    return [r for r in all_rows if _coerce_str(r.get('company_name'))]


def _llm_extract_signals(text):
    """Use LLM to extract intent signals from unstructured document text.

    Returns a list of dicts with canonical keys, or None on failure.
    """
    system_prompt = (
        "You are an assistant that extracts sales intent signals from unstructured text. "
        "The user works at Phrase, a localization and translation management platform. "
        "Extract every company/signal pair mentioned. For each, output a JSON array "
        "where each element has these keys:\n"
        "- company_name (required)\n"
        "- signal_description (required — what the signal is, in 1-2 sentences)\n"
        "- signal_type (optional — e.g. 'hiring_localization', 'new_market_expansion', "
        "'i18n_library_adoption', 'website_translation', 'tms_evaluation')\n"
        "- website (optional — company domain if mentioned)\n"
        "- industry (optional)\n"
        "- account_owner (optional — if an owner/rep is mentioned)\n\n"
        "Output ONLY valid JSON. No explanation. If no signals found, output []."
    )

    # Truncate very long documents
    truncated = text[:8000]
    user_prompt = f"Extract intent signals from this document:\n\n{truncated}"

    response = _llm_generate(system_prompt, user_prompt)
    if not response:
        return None

    try:
        cleaned = response.strip()
        # Strip markdown code fences if present
        if cleaned.startswith('```'):
            cleaned = cleaned.split('\n', 1)[1].rsplit('```', 1)[0]
        parsed = json.loads(cleaned)
        if isinstance(parsed, list) and len(parsed) > 0:
            return parsed
    except (json.JSONDecodeError, IndexError, ValueError):
        logger.warning("[INGEST] Could not parse LLM signal extraction response")

    return None


# ---------------------------------------------------------------------------
# Clear All Signals (for "Replace all" imports)
# ---------------------------------------------------------------------------

def _clear_all_signals():
    """Delete all intent_signals and monitored_accounts rows.

    Returns the number of signals deleted.
    """
    with db_connection() as conn:
        cursor = conn.cursor()
        # Delete signals first (FK dependency)
        cursor.execute("SELECT COUNT(*) as cnt FROM intent_signals")
        row = cursor.fetchone()
        count = row['cnt'] if isinstance(row, dict) else row[0]

        cursor.execute("DELETE FROM drafts")
        cursor.execute("DELETE FROM prospects")
        cursor.execute("DELETE FROM intent_signals")
        cursor.execute("DELETE FROM monitored_accounts")
        conn.commit()

        logger.info("[INGEST] Cleared all signals and accounts (%d signals deleted)", count)
        return count


# ---------------------------------------------------------------------------
# Post-Processing: Apollo Enrichment + BDR Evaluation
# ---------------------------------------------------------------------------

def _post_process_batch(new_accounts, created_signals):
    """Run Apollo enrichment on new accounts and BDR evaluation on signals.

    Both are best-effort — failures don't affect the import result.
    Returns a summary dict of what was enriched/evaluated.
    """
    result = {
        'accounts_enriched': 0,
        'signals_evaluated': 0,
    }

    if not new_accounts and not created_signals:
        return result

    # 1. Apollo enrichment for new accounts
    result['accounts_enriched'] = _enrich_accounts_apollo(new_accounts)

    # 2. BDR evaluation for all created signals
    result['signals_evaluated'] = _evaluate_signals_bdr(created_signals)

    # 3. Consolidate signals per account (merge multiple signals into one)
    result['signals_consolidated'] = _consolidate_batch_signals(created_signals)

    return result


def _consolidate_batch_signals(created_signals):
    """Consolidate signals per account after a batch import.

    If an account ended up with multiple active signals (from this batch or
    combined with pre-existing ones), merge them into one consolidated signal.

    Returns the number of accounts consolidated.
    """
    if not created_signals:
        return 0

    try:
        from v2.services.consolidation_service import consolidate_account
        from v2.db import db_connection, rows_to_dicts

        # Find unique account_ids from this batch that have 2+ active signals
        seen_accounts = set()
        for sig in created_signals:
            # Get account_id from the signal record
            with db_connection() as conn:
                cursor = conn.cursor()
                cursor.execute(
                    "SELECT account_id FROM intent_signals WHERE id = ?",
                    (sig.get('signal_id'),),
                )
                row = cursor.fetchone()
                if row:
                    aid = row['account_id'] if isinstance(row, dict) else row[0]
                    seen_accounts.add(aid)

        consolidated = 0
        for account_id in seen_accounts:
            result = consolidate_account(account_id)
            if result:
                consolidated += 1

        if consolidated:
            logger.info("[INGEST] Consolidated signals for %d accounts", consolidated)
        return consolidated

    except Exception as e:
        logger.warning("[INGEST] Signal consolidation failed (non-fatal): %s", e)
        return 0


def _enrich_accounts_apollo(accounts):
    """Look up new accounts on Apollo to fill missing fields (industry, size, etc.).

    Returns the number of accounts successfully enriched.
    """
    if not accounts:
        return 0

    apollo_key = os.environ.get('APOLLO_API_KEY', '')
    if not apollo_key:
        logger.info("[ENRICH] Apollo API key not set — skipping enrichment")
        return 0

    try:
        from apollo_client import apollo_api_call
    except ImportError:
        logger.info("[ENRICH] apollo_client not available — skipping enrichment")
        return 0

    enriched_count = 0

    for acct in accounts:
        try:
            org = None
            website = acct.get('website') or ''

            if website:
                # Strip protocol to get domain
                domain = website.replace('https://', '').replace('http://', '').split('/')[0]
                if domain:
                    resp = apollo_api_call(
                        'get',
                        f'https://api.apollo.io/api/v1/organizations/enrich?domain={domain}',
                    )
                    if resp.status_code == 200:
                        org = resp.json().get('organization')

            if not org:
                # Fall back to name search
                resp = apollo_api_call(
                    'post',
                    'https://api.apollo.io/v1/mixed_companies/search',
                    json={
                        'q_organization_name': acct['company_name'],
                        'per_page': 1,
                    },
                )
                if resp.status_code == 200:
                    orgs = resp.json().get('organizations', [])
                    org = orgs[0] if orgs else None

            if not org:
                continue

            # Map Apollo fields to our account fields
            updates = {}
            if org.get('website_url'):
                updates['website'] = org['website_url']
            if org.get('industry'):
                updates['industry'] = org['industry']
            if org.get('estimated_num_employees'):
                updates['employee_count'] = str(org['estimated_num_employees'])
                # Also set company_size from employee count ranges
                emp = org['estimated_num_employees']
                if emp < 50:
                    updates['company_size'] = '1-50'
                elif emp < 200:
                    updates['company_size'] = '51-200'
                elif emp < 1000:
                    updates['company_size'] = '201-1000'
                elif emp < 5000:
                    updates['company_size'] = '1001-5000'
                else:
                    updates['company_size'] = '5000+'
            if org.get('annual_revenue_printed'):
                updates['annual_revenue'] = org['annual_revenue_printed']
            if org.get('linkedin_url'):
                updates['linkedin_url'] = org['linkedin_url']

            # Build HQ location from Apollo data
            hq_parts = []
            if org.get('city'):
                hq_parts.append(org['city'])
            if org.get('state'):
                hq_parts.append(org['state'])
            if org.get('country'):
                hq_parts.append(org['country'])
            if hq_parts:
                updates['hq_location'] = ', '.join(hq_parts)

            if org.get('funding_stage'):
                updates['funding_stage'] = org['funding_stage']

            # Resolve real company name if current name looks like a GitHub org login
            # (short, no spaces = likely an org slug like "gf" instead of "General Fasteners")
            current_name = acct.get('company_name', '')
            apollo_name = (org.get('name') or '').strip()
            if (apollo_name
                    and current_name
                    and ' ' not in current_name
                    and len(current_name) <= 20
                    and apollo_name.lower() != current_name.lower()):
                updates['company_name'] = apollo_name
                logger.info("[ENRICH] Resolved company name: %s → %s",
                            current_name, apollo_name)

            if updates:
                account_service.update_account_enrichment(acct['account_id'], **updates)
                enriched_count += 1
                logger.info("[ENRICH] Enriched %s with %d fields from Apollo",
                            acct['company_name'], len(updates))

        except Exception as e:
            logger.warning("[ENRICH] Apollo enrichment failed for %s: %s",
                           acct['company_name'], e)

    return enriched_count


def _evaluate_signals_bdr(signals_data):
    """LLM evaluates signal quality and positioning for a batch of signals.

    Sends all signals in one LLM call, parses the response, and updates
    each signal's bdr_quality_score and bdr_positioning.

    Returns the number of signals successfully evaluated.
    """
    if not signals_data:
        return 0

    if not _get_llm_client():
        logger.info("[EVAL] LLM not available — skipping BDR evaluation")
        return 0

    system_prompt = (
        "You are a senior BDR at Phrase, a localization and translation management platform "
        "(phrase.com). Evaluate each intent signal for cold outreach potential.\n\n"
        "For each signal, provide:\n"
        "1. quality_score (1-5):\n"
        "   5 = Strong buying signal, clear pain point, urgent need for localization\n"
        "   4 = Good signal, clear positioning opportunity for Phrase\n"
        "   3 = Moderate signal, needs more context but worth pursuing\n"
        "   2 = Weak signal, generic or unclear relevance to localization\n"
        "   1 = Not useful for outreach\n"
        "2. positioning: A concise 1-2 sentence angle for how to use this signal "
        "in cold outreach email. Be specific about what Phrase offers that helps.\n\n"
        "Output ONLY a JSON array with objects containing: "
        "signal_id (int), quality_score (int 1-5), positioning (string).\n"
        "No explanation, just JSON."
    )

    # Build the signals summary for the LLM
    signals_for_llm = [
        {
            'signal_id': s['signal_id'],
            'company': s['company_name'],
            'description': s['signal_description'],
            'type': s.get('signal_type') or 'unknown',
        }
        for s in signals_data
    ]

    user_prompt = f"Evaluate these {len(signals_for_llm)} intent signals:\n\n{json.dumps(signals_for_llm, indent=2)}"

    response = _llm_generate(system_prompt, user_prompt)
    if not response:
        return 0

    try:
        cleaned = response.strip()
        if cleaned.startswith('```'):
            cleaned = cleaned.split('\n', 1)[1].rsplit('```', 1)[0]
        evaluations = json.loads(cleaned)
    except (json.JSONDecodeError, IndexError, ValueError):
        logger.warning("[EVAL] Could not parse LLM BDR evaluation response")
        return 0

    evaluated_count = 0
    for ev in evaluations:
        try:
            sid = ev.get('signal_id')
            score = int(ev.get('quality_score', 0))
            positioning = ev.get('positioning', '')
            if sid and 1 <= score <= 5:
                signal_service.update_signal_bdr_evaluation(sid, score, positioning)
                evaluated_count += 1
        except Exception:
            pass

    logger.info("[EVAL] BDR evaluation complete: %d/%d signals evaluated",
                evaluated_count, len(signals_data))
    return evaluated_count


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_column(headers, headers_lower, candidates):
    """Find the first matching header from a tuple of candidate names.

    Returns the original-case header name, or None if no match.
    """
    for candidate in candidates:
        for orig, low in zip(headers, headers_lower):
            if low == candidate:
                return orig
    return None
